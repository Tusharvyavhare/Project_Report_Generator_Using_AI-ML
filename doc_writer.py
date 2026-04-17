from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import re

def add_formatted_text(para, text):
    """Add markdown-style formatted text (**bold**, *italic*)"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        else:
            subparts = re.split(r'(\*.*?\*)', part)
            for subpart in subparts:
                if subpart.startswith('*') and subpart.endswith('*') and len(subpart) > 2:
                    run = para.add_run(subpart[1:-1])
                    run.italic = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                else:
                    run = para.add_run(subpart)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

def clear_body_paragraphs(doc):
    """Remove all paragraphs from body but keep header/footer"""
    for para in doc.paragraphs[:]:
        p = para._element
        p.getparent().remove(p)

def _create_field_run(para, field_code):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = para.add_run()
    fldCharBegin = OxmlElement('w:fldChar')
    fldCharBegin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldCharBegin)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_code
    run._r.append(instrText)

    fldCharSeparate = OxmlElement('w:fldChar')
    fldCharSeparate.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldCharSeparate)

    run_placeholder = para.add_run('1')
    fldCharEnd = OxmlElement('w:fldChar')
    fldCharEnd.set(qn('w:fldCharType'), 'end')
    run_placeholder._r.append(fldCharEnd)

def _set_page_border(doc):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    for section in doc.sections:
        sectPr = section._sectPr
        borders = sectPr.find(qn('w:pgBorders'))
        if borders is None:
            borders = OxmlElement('w:pgBorders')
            sectPr.append(borders)
        for side in ('top', 'left', 'bottom', 'right'):
            border = borders.find(qn(f'w:{side}'))
            if border is None:
                border = OxmlElement(f'w:{side}')
                borders.append(border)
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), '000000')

def _apply_common_report_formatting(doc, project_title):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Header
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.text = project_title
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Footer with page numbers
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.text = 'Page '
    _create_field_run(footer_para, 'PAGE')

    # Page border
    _set_page_border(doc)

def rebuild_doc(*, project_title, structure, generated_sections, output_path, template_path=None):
    """Create a new Word report, remove old body content, apply formatting."""
    doc = Document(template_path) if template_path else Document()
    clear_body_paragraphs(doc)

    # Main title
    title_para = doc.add_paragraph(project_title)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)

    # Add sections
    sec_idx = 0
    for item in structure:
        if item.get("heading_level"):
            heading_para = doc.add_paragraph(item.get("text", ""))
            heading_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in heading_para.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)

            if sec_idx < len(generated_sections):
                content_para = doc.add_paragraph()
                add_formatted_text(content_para, generated_sections[sec_idx])
                content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                sec_idx += 1

    # Remaining sections
    while sec_idx < len(generated_sections):
        content_para = doc.add_paragraph()
        add_formatted_text(content_para, generated_sections[sec_idx])
        content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        sec_idx += 1

    # Header/footer/border
    _apply_common_report_formatting(doc, project_title)
    doc.save(output_path)